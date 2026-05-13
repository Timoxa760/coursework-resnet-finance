"""Сборка курсовой работы в формате DOCX строго по ГОСТ.

Требования (по методичке Финуниверситета):
- Шрифт: Times New Roman, 14 пт
- Интервал: 1.5
- Поля: левое 30 мм, правое 15 мм, верхнее 20 мм, нижнее 20 мм
- Абзацный отступ: 1.25 см
- Заголовки разделов: полужирный, прописные, без отступа
- Заголовки подразделов: полужирный, без отступа
- Формулы: выравнивание по центру, номер справа
- Таблицы: границы, название сверху слева
- Рисунки: подпись снизу по центру
"""

import logging
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, Inches, Mm
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

FILES = [
    "01_title.md",
    "02_intro.md",
    "03_chapter1.md",
    "04_chapter2.md",
    "05_conclusion.md",
    "06_references.md",
    "07_appendix.md",
]


def set_run_font(run, font_name="Times New Roman", font_size=14, bold=False):
    """Устанавливает шрифт для run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_format(
    paragraph,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Cm(1.25),
    line_spacing=1.5,
    space_after=Pt(0),
    space_before=Pt(0),
):
    """Устанавливает формат параграфа по ГОСТ."""
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.first_line_indent = first_line_indent
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = space_before
    pf.widow_control = True


def apply_gost_styles(doc: Document):
    """Применяет стили ГОСТ ко всем элементам документа."""
    logger.info("Применение стилей ГОСТ...")

    # Настройка полей страницы
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)

    # Стили заголовков
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(14)
            style.font.bold = True
            style.font.color.rgb = None
            pf = style.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            pf.keep_with_next = True
            pf.keep_together = True
        except KeyError:
            logger.warning("Стиль %s не найден", style_name)

    # Стиль Normal
    try:
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = False
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.widow_control = True
    except KeyError:
        logger.warning("Стиль Normal не найден")

    # Обработка всех параграфов
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name

        # Определяем, является ли параграф заголовком
        is_heading = style_name.startswith("Heading")
        is_title = style_name == "Title"

        for run in paragraph.runs:
            set_run_font(
                run,
                font_name="Times New Roman",
                font_size=14,
                bold=is_heading or is_title,
            )

        pf = paragraph.paragraph_format
        if is_heading:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.line_spacing = 1.5
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            pf.keep_with_next = True
            pf.keep_together = True
            # Заголовки разделов — прописные (для Heading 1)
            if style_name == "Heading 1" and paragraph.text:
                paragraph.text = paragraph.text.upper()
        elif is_title:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
        else:
            # Обычный текст
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Cm(1.25)
            pf.line_spacing = 1.5
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)

    # Обработка таблиц
    for table in doc.tables:
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "Times New Roman", 12, False)
                    pf = paragraph.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pf.first_line_indent = Cm(0)
                    pf.line_spacing = 1.0
                    pf.space_after = Pt(0)
                    pf.space_before = Pt(0)

    logger.info("Стили ГОСТ применены")


def merge_markdown(text_dir: Path, output: Path) -> None:
    """Объединяет markdown-файлы в один."""
    with open(output, "w", encoding="utf-8") as out:
        for fname in FILES:
            fpath = text_dir / fname
            if not fpath.exists():
                logger.warning("Файл не найден: %s", fpath)
                continue
            with open(fpath, "r", encoding="utf-8") as f:
                out.write(f.read())
                out.write("\n\n")
    logger.info("Объединённый Markdown сохранён: %s", output)


def build_docx(draft: Path, output: Path) -> None:
    """Конвертирует Markdown в DOCX через pandoc."""
    pandoc_path = r"C:\Windows\Temp\pandoc\pandoc-3.1.11\pandoc.exe"
    cmd = [
        pandoc_path,
        str(draft),
        "-o",
        str(output),
        "--from",
        "markdown",
        "--to",
        "docx",
        "--toc",
        "--toc-depth=2",
        "--highlight-style=tango",
        "-V",
        "lang=ru",
    ]
    logger.info("Запуск pandoc: %s", " ".join(cmd))
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        logger.error("Pandoc ошибка: %s", result.stderr)
        raise RuntimeError(result.stderr)
    logger.info("Базовый DOCX собран: %s", output)


def main():
    logging.basicConfig(level=logging.INFO)
    text_dir = Path(__file__).parent
    draft = text_dir / "draft.md"
    raw_docx = text_dir / "coursework_raw.docx"
    final_docx = text_dir.parent / "Миндрин_Тимофей_Курсовая.docx"

    merge_markdown(text_dir, draft)
    build_docx(draft, raw_docx)

    logger.info("Применение стилей ГОСТ...")
    doc = Document(raw_docx)
    apply_gost_styles(doc)
    # Сохраняем во временный файл, затем копируем
    temp_final = text_dir / "coursework_temp.docx"
    doc.save(temp_final)
    logger.info("Временный DOCX сохранён: %s", temp_final)

    # Копируем с заменой
    import shutil
    shutil.copy(str(temp_final), str(final_docx))
    logger.info("Итоговый DOCX сохранён: %s", final_docx)

    # Удаляем временные файлы
    raw_docx.unlink(missing_ok=True)
    temp_final.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
